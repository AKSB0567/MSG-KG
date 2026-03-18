"""
EMNLP Evaluation Pipeline — Tables 4 & 5
==========================================
Evidence-Grounded Reasoning Performance (Table 4) and Ablation Study (Table 5).

Generator: Qwen2.5-7B-Instruct (Ollama, local GPU)
Judge:     Llama-3.1-8B-Instruct (Ollama, local GPU) — different model family

Protocol:
  - 273 test questions (3 per company × 91 companies)
  - 5 reasoning systems + 3 ablation variants
  - LLM-as-Judge with explicit rubrics (separate model)
  - Bootstrap 95% CIs, paired significance tests

Run:  python eval_reasoning_pipeline.py
"""

import json, re, pathlib, time, random, sys, os
import urllib.request
import numpy as np
from collections import defaultdict

# ── Config ──────────────────────────────────────────────────────────────
BASE_DIR    = pathlib.Path(__file__).parent
REG_PATH    = BASE_DIR / "companies_registry.json"
SEC_DIR     = BASE_DIR / "MSGKG" / "data" / "data"
ONTO_DIR    = BASE_DIR / "data" / "mission_kg" / "ontology"
RESULTS_DIR = BASE_DIR / "data" / "eval_results"
RESULTS_DIR.mkdir(parents=True, exist_ok=True)

OLLAMA_URL      = "http://localhost:11434"
GENERATOR_MODEL = "qwen2.5:7b"
JUDGE_MODEL     = "llama3.1:8b"

N_BOOTSTRAP     = 1000
CONFIDENCE      = 0.95
HUMAN_EVAL_SIZE = 50

# ── Ollama helpers ──────────────────────────────────────────────────────

def ollama_chat(model, messages, max_tokens=400, json_mode=False):
    """Call Ollama chat API. Returns response text."""
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {"temperature": 0.1, "num_predict": max_tokens},
    }
    if json_mode:
        payload["format"] = "json"
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/chat", data=data,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result["message"]["content"].strip()
    except Exception as e:
        print(f"    LLM error: {str(e)[:60]}")
        return ""


def ollama_generate(model, prompt, max_tokens=20):
    """Simple generate call (for model loading/unloading)."""
    payload = json.dumps({
        "model": model, "prompt": prompt, "stream": False,
        "options": {"num_predict": max_tokens},
    }).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/generate", data=payload,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except:
        return {}


def swap_to_model(target_model, other_model):
    """Unload other model, load target model."""
    print(f"  Swapping to {target_model}...")
    # Unload other
    payload = json.dumps({
        "model": other_model, "prompt": "", "stream": False,
        "keep_alive": 0,
    }).encode("utf-8")
    try:
        req = urllib.request.Request(
            f"{OLLAMA_URL}/api/generate", data=payload,
            headers={"Content-Type": "application/json"},
        )
        urllib.request.urlopen(req, timeout=30)
    except:
        pass
    time.sleep(2)
    # Load target
    ollama_generate(target_model, "hi", max_tokens=5)
    print(f"  {target_model} loaded.")


def parse_json_response(text):
    """Extract JSON from LLM response."""
    start = text.find("{")
    end = text.rfind("}") + 1
    if start >= 0 and end > start:
        try:
            return json.loads(text[start:end])
        except json.JSONDecodeError:
            pass
    return None


# ── Data loading ────────────────────────────────────────────────────────

def load_registry():
    with open(REG_PATH, encoding="utf-8") as f:
        return json.load(f)


def extract_item1(filepath):
    """Extract Item 1 (Business) from 10-K."""
    text = filepath.read_text(encoding="utf-8", errors="ignore")
    hdr = text.find("</SEC-HEADER>")
    if hdr > 0:
        text = text[hdr + len("</SEC-HEADER>"):]
    text = re.sub(r'\s+', ' ', text)
    m = re.search(r'(?:ITEM\s*1\.?\s*(?:BUSINESS|Business))', text, re.IGNORECASE)
    start = m.end() if m else 0
    em = re.search(r'(?:ITEM\s*1A\.?\s|ITEM\s*2\.?\s|PART\s*II\b)', text[start:], re.IGNORECASE)
    end = start + em.start() if em else min(start + 200000, len(text))
    return text[start:end].strip()[:100000]


def find_10k_file(cik):
    cik_stripped = cik.lstrip("0")
    for f in SEC_DIR.glob("cleaned_10-K_*.txt"):
        acc = f.stem.replace("cleaned_10-K_", "")
        file_cik = acc.split("-")[0].lstrip("0")
        if file_cik == cik_stripped:
            return f
    return None


def chunk_text(text, chunk_size=500, overlap=80):
    """Split text into word-based chunks."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        if len(chunk_words) < 30:
            break
        chunks.append(" ".join(chunk_words))
        i += chunk_size - overlap
    return chunks


def slug(name):
    """Convert company name to safe filename."""
    import unicodedata
    s = unicodedata.normalize('NFKD', name).encode('ascii', 'ignore').decode()
    s = re.sub(r'[^\w\s-]', '', s).strip().lower()
    return re.sub(r'[-\s]+', '_', s)[:50]


def parse_ontology_ttl(cik, company_name):
    """Parse ontology TTL and extract mission-strategy structure."""
    ttl_path = ONTO_DIR / f"{slug(company_name)}.ttl"
    if not ttl_path.exists():
        return {"mission": "", "stakeholders": [], "values": [], "objectives": [], "capabilities": [], "domains": []}

    content = ttl_path.read_text(encoding="utf-8", errors="ignore")
    result = {"mission": "", "stakeholders": [], "values": [], "objectives": [], "capabilities": [], "domains": []}

    # Extract mission text
    m = re.search(r'msg:missionText\s+"([^"]*)"', content)
    if m:
        result["mission"] = m.group(1)

    # Extract entities by type
    type_map = {
        "Stakeholder": "stakeholders",
        "CorporateValue": "values",
        "CorporateObjective": "objectives",
        "BusinessCapability": "capabilities",
        "BusinessDomain": "domains",
    }
    for line in content.split("\n"):
        for rdf_type, key in type_map.items():
            if f"a msg:{rdf_type}" in line or f"a fibo-be:{rdf_type}" in line:
                # Get label from next lines
                inst_match = re.match(r'(inst:\S+)', line)
                if inst_match:
                    inst_name = inst_match.group(1).replace("inst:", "")
                    # Clean up name
                    clean = re.sub(r'_', ' ', inst_name.split("_", 1)[-1] if "_" in inst_name else inst_name)
                    result[key].append(clean)

    # Also grab labels
    for m in re.finditer(r'rdfs:label\s+"([^"]*)"', content):
        label = m.group(1)
        # Heuristic: check context before this label
        pos = m.start()
        context = content[max(0, pos - 200):pos]
        for rdf_type, key in type_map.items():
            if rdf_type in context:
                if label not in result[key]:
                    result[key].append(label)
                break

    return result


# ── Question generation ─────────────────────────────────────────────────

QUESTION_TEMPLATES = [
    "What is {company}'s stated mission, and what specific evidence from their 10-K filing supports it?",
    "How does {company}'s mission connect to their strategic priorities and operational capabilities?",
    "Trace how {company}'s mission is operationalized through specific strategic initiatives and mechanisms described in their 10-K filing.",
]


def generate_questions(registry):
    """Generate 3 questions per company = 273 total."""
    questions = []
    for cik, info in registry.items():
        name = info["name"]
        for i, template in enumerate(QUESTION_TEMPLATES):
            questions.append({
                "qid": f"{cik}_q{i+1}",
                "cik": cik,
                "company": name,
                "sector": info.get("sector", "Unknown"),
                "question": template.format(company=name),
                "q_type": ["factual", "strategic", "trace"][i],
            })
    return questions


# ── Retrieval engines ───────────────────────────────────────────────────

def tfidf_retrieve(query, chunks, top_k=5):
    """TF-IDF based retrieval."""
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity

    if not chunks:
        return []
    corpus = chunks + [query]
    vectorizer = TfidfVectorizer(max_features=1000, stop_words="english")
    tfidf = vectorizer.fit_transform(corpus)
    query_vec = tfidf[-1]
    doc_vecs = tfidf[:-1]
    sims = cosine_similarity(query_vec, doc_vecs).flatten()
    top_idx = sims.argsort()[::-1][:top_k]
    return [chunks[i] for i in top_idx if sims[i] > 0]


_embedder = None
def get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer
        _embedder = SentenceTransformer("all-MiniLM-L6-v2")
    return _embedder


def semantic_retrieve(query, chunks, top_k=5):
    """Sentence-transformer based semantic retrieval."""
    if not chunks:
        return []
    model = get_embedder()
    query_emb = model.encode([query], normalize_embeddings=True)
    chunk_embs = model.encode(chunks, normalize_embeddings=True, show_progress_bar=False)
    sims = np.dot(chunk_embs, query_emb.T).flatten()
    top_idx = sims.argsort()[::-1][:top_k]
    return [chunks[i] for i in top_idx if sims[i] > 0]


# ── 5 Reasoning Systems ────────────────────────────────────────────────

def system_text_only(question, **kwargs):
    """System 1: Text-Only LLM — no retrieval context."""
    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer questions about corporate strategy based on your knowledge."},
        {"role": "user", "content": question},
    ], max_tokens=300)
    return resp


def system_vector_retrieval(question, chunks, **kwargs):
    """System 2: Vector Retrieval (TF-IDF) + LLM."""
    retrieved = tfidf_retrieve(question, chunks, top_k=5)
    context = "\n\n".join(retrieved[:5])
    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based ONLY on the provided evidence passages. Cite specific passages."},
        {"role": "user", "content": f"Evidence passages from 10-K filing:\n{context[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=300)
    return resp


def system_rag(question, chunks, **kwargs):
    """System 3: SentenceTransformer + RAG."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    context = "\n\n".join(retrieved[:5])
    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based ONLY on the provided evidence passages. Cite specific passages."},
        {"role": "user", "content": f"Retrieved passages from 10-K filing:\n{context[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=300)
    return resp


def system_graphrag(question, chunks, kg_struct, **kwargs):
    """System 4: GraphRAG — KG triples + text retrieval."""
    # Build KG context from ontology structure
    kg_lines = []
    if kg_struct.get("mission"):
        kg_lines.append(f"Mission: {kg_struct['mission']}")
    for s in kg_struct.get("stakeholders", []):
        kg_lines.append(f"Mission → serves → {s}")
    for v in kg_struct.get("values", []):
        kg_lines.append(f"Mission → embodies → {v}")
    for o in kg_struct.get("objectives", []):
        kg_lines.append(f"Mission → pursues → {o}")
    for c in kg_struct.get("capabilities", []):
        kg_lines.append(f"Strategy → leverages → {c}")
    for d in kg_struct.get("domains", []):
        kg_lines.append(f"Company → operates in → {d}")
    kg_context = "\n".join(kg_lines) if kg_lines else "No knowledge graph available."

    # Also retrieve text
    retrieved = semantic_retrieve(question, chunks, top_k=3)
    text_context = "\n\n".join(retrieved[:3])

    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Use the knowledge graph triples AND text evidence to answer. Cite both graph relationships and textual evidence."},
        {"role": "user", "content": f"Knowledge Graph:\n{kg_context}\n\nText Evidence:\n{text_context[:2000]}\n\nQuestion: {question}"},
    ], max_tokens=400)
    return resp


def system_msgkg(question, chunks, kg_struct, mission, evidence_spans=None, **kwargs):
    """System 5: MSG-KG — Full pipeline with reasoning paths + evidence."""
    # Build structured reasoning path
    path_lines = []
    if mission:
        path_lines.append(f"MISSION: {mission}")
    pillars = kg_struct.get("values", []) + kg_struct.get("objectives", [])
    if pillars:
        path_lines.append(f"STRATEGIC PILLARS: {', '.join(pillars[:5])}")
        for p in pillars[:3]:
            path_lines.append(f"  Mission → {p}")
    mechanisms = kg_struct.get("capabilities", [])
    if mechanisms:
        path_lines.append(f"OPERATIONAL MECHANISMS: {', '.join(mechanisms[:5])}")
        for m in mechanisms[:3]:
            # Link to pillars
            if pillars:
                path_lines.append(f"  {pillars[0]} → {m}")
    reasoning_path = "\n".join(path_lines) if path_lines else "No reasoning path available."

    # Evidence spans from 10-K
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(f"[Evidence {i+1}]: {r[:300]}" for i, r in enumerate(retrieved[:5]))

    # KG triples (full)
    kg_lines = []
    for s in kg_struct.get("stakeholders", []):
        kg_lines.append(f"serves_stakeholder({s})")
    for v in kg_struct.get("values", []):
        kg_lines.append(f"embodies_value({v})")
    for o in kg_struct.get("objectives", []):
        kg_lines.append(f"pursues_objective({o})")
    for c in kg_struct.get("capabilities", []):
        kg_lines.append(f"leverages_capability({c})")
    kg_triples = "; ".join(kg_lines) if kg_lines else "N/A"

    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst performing evidence-grounded reasoning. Use the structured reasoning path, knowledge graph, AND textual evidence to answer. Explicitly trace the Mission → Strategic Pillar → Operational Mechanism chain. Cite specific evidence."},
        {"role": "user", "content": (
            f"Mission-Strategy Reasoning Path:\n{reasoning_path}\n\n"
            f"KG Relations: {kg_triples}\n\n"
            f"Evidence from 10-K Filing:\n{evidence_text[:2500]}\n\n"
            f"Question: {question}\n\n"
            f"Answer with explicit reasoning chain (Mission → Pillar → Mechanism) and cite evidence:"
        )},
    ], max_tokens=500)
    return resp


# ── Ablation variants ───────────────────────────────────────────────────

def ablation_no_graph(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without Graph Retrieval — text retrieval only."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(f"[Evidence {i+1}]: {r[:300]}" for i, r in enumerate(retrieved[:5]))

    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based on the mission statement and textual evidence. Cite evidence."},
        {"role": "user", "content": f"Mission: {mission}\n\nEvidence:\n{evidence_text[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=400)
    return resp


def ablation_no_evidence(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without Evidence Linking — KG path only, no text evidence."""
    path_lines = [f"MISSION: {mission}"]
    pillars = kg_struct.get("values", []) + kg_struct.get("objectives", [])
    mechanisms = kg_struct.get("capabilities", [])
    for p in pillars[:5]:
        path_lines.append(f"  Mission → {p}")
    for m in mechanisms[:5]:
        if pillars:
            path_lines.append(f"  {pillars[0]} → {m}")
    reasoning_path = "\n".join(path_lines)

    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Use the structured reasoning path to answer. Trace the Mission → Pillar → Mechanism chain."},
        {"role": "user", "content": f"Reasoning Path:\n{reasoning_path}\n\nQuestion: {question}"},
    ], max_tokens=400)
    return resp


def ablation_no_kg_structure(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without KG Structure — flat retrieval, no structured path."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(retrieved[:5])

    # Flat list of KG elements (no structure)
    elements = []
    for key in ["stakeholders", "values", "objectives", "capabilities", "domains"]:
        elements.extend(kg_struct.get(key, []))
    flat_kg = ", ".join(elements) if elements else "N/A"

    resp = ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based on the retrieved text and related concepts. Cite evidence."},
        {"role": "user", "content": f"Mission: {mission}\nRelated concepts: {flat_kg}\n\nText:\n{evidence_text[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=400)
    return resp


# ── LLM-as-Judge ────────────────────────────────────────────────────────

GROUNDING_RUBRIC = """Score the answer's evidence grounding on a scale of 1-5:
5: All claims are supported by specific, verifiable evidence from the filing
4: Most claims supported, minor unsupported assertions
3: Some evidence cited but significant claims lack support
2: Minimal evidence, mostly unsupported claims
1: No evidence grounding, entirely generic or hallucinated

COMPANY: {company}
SECTOR: {sector}
REFERENCE MISSION: {mission}

QUESTION: {question}
ANSWER TO EVALUATE: {answer}

Respond with ONLY JSON: {{"grounding_score": N, "reason": "one sentence"}}"""

EXPLANATION_RUBRIC = """Score the answer's structural reasoning quality on a scale of 1-5:
5: Clear Mission → Strategic Pillar → Operational Mechanism chain with specific evidence
4: Mostly complete chain, minor gaps in reasoning
3: Partial chain, some strategic elements connected but incomplete
2: Weak structural reasoning, mostly flat description
1: No structural reasoning, just generic statements

COMPANY: {company}
COMPANY MISSION: {mission}
KNOWN STRATEGIC ELEMENTS: {strategic_elements}

QUESTION: {question}
ANSWER TO EVALUATE: {answer}

Respond with ONLY JSON: {{"explanation_score": N, "reason": "one sentence"}}"""


def judge_answer(question_data, answer, mission, kg_struct):
    """Judge a single answer using Llama-3.1-8B. Returns (grounding_score, explanation_score)."""
    company = question_data["company"]
    sector = question_data["sector"]
    question = question_data["question"]

    strategic_elements = ", ".join(
        kg_struct.get("values", [])[:3] + kg_struct.get("objectives", [])[:3] +
        kg_struct.get("capabilities", [])[:3]
    )

    # Grounding score
    g_resp = ollama_chat(JUDGE_MODEL, [
        {"role": "system", "content": "You are an expert evaluator. Score answers strictly. Respond with JSON only."},
        {"role": "user", "content": GROUNDING_RUBRIC.format(
            company=company, sector=sector, mission=mission[:200],
            question=question, answer=answer[:500],
        )},
    ], max_tokens=100, json_mode=True)

    g_json = parse_json_response(g_resp)
    g_score = g_json.get("grounding_score", 2) if g_json else 2

    # Explanation score
    e_resp = ollama_chat(JUDGE_MODEL, [
        {"role": "system", "content": "You are an expert evaluator. Score answers strictly. Respond with JSON only."},
        {"role": "user", "content": EXPLANATION_RUBRIC.format(
            company=company, mission=mission[:200],
            strategic_elements=strategic_elements[:300],
            question=question, answer=answer[:500],
        )},
    ], max_tokens=100, json_mode=True)

    e_json = parse_json_response(e_resp)
    e_score = e_json.get("explanation_score", 2) if e_json else 2

    # Clamp to 1-5
    g_score = max(1, min(5, int(g_score)))
    e_score = max(1, min(5, int(e_score)))

    return g_score, e_score


# ── Statistical analysis ────────────────────────────────────────────────

def normalize_score(raw_scores):
    """Convert 1-5 Likert to 0-100 scale."""
    return [(s - 1) / 4 * 100 for s in raw_scores]


def bootstrap_ci(scores, n_bootstrap=N_BOOTSTRAP, ci=CONFIDENCE):
    """Bootstrap 95% confidence interval. Returns (mean, lower, upper)."""
    scores = np.array(scores)
    means = []
    for _ in range(n_bootstrap):
        sample = np.random.choice(scores, size=len(scores), replace=True)
        means.append(np.mean(sample))
    means = sorted(means)
    alpha = (1 - ci) / 2
    lower = means[int(alpha * n_bootstrap)]
    upper = means[int((1 - alpha) * n_bootstrap)]
    return float(np.mean(scores)), float(lower), float(upper)


def paired_bootstrap_test(scores_a, scores_b, n_bootstrap=N_BOOTSTRAP):
    """Paired bootstrap test. Returns p-value for H0: mean(A) <= mean(B)."""
    scores_a = np.array(scores_a)
    scores_b = np.array(scores_b)
    observed_diff = np.mean(scores_a) - np.mean(scores_b)
    count = 0
    for _ in range(n_bootstrap):
        idx = np.random.choice(len(scores_a), size=len(scores_a), replace=True)
        diff = np.mean(scores_a[idx]) - np.mean(scores_b[idx])
        if diff <= 0:
            count += 1
    return count / n_bootstrap


# ── Human evaluation subset ─────────────────────────────────────────────

def select_human_eval_subset(all_results, n=HUMAN_EVAL_SIZE):
    """Select stratified subset for human evaluation."""
    # Stratify by question type and score variance
    by_type = defaultdict(list)
    for r in all_results:
        by_type[r["q_type"]].append(r)

    selected = []
    per_type = n // 3
    for q_type, items in by_type.items():
        # Sort by score variance across systems (most interesting first)
        for item in items:
            scores = [item["scores"].get(sys, {}).get("grounding", 3) for sys in item["scores"]]
            item["_variance"] = np.var(scores) if scores else 0
        items.sort(key=lambda x: -x["_variance"])
        selected.extend(items[:per_type])

    # Fill remaining
    remaining = [r for r in all_results if r not in selected]
    random.shuffle(remaining)
    selected.extend(remaining[:n - len(selected)])

    return selected[:n]


# ── Main pipeline ───────────────────────────────────────────────────────

def main():
    print("=" * 70)
    print("EMNLP Evaluation Pipeline — Tables 4 & 5")
    print("Generator: Qwen2.5-7B | Judge: Llama-3.1-8B")
    print("=" * 70)

    # ── Phase 0: Load data ──
    print("\n[Phase 0] Loading data...")
    registry = load_registry()
    print(f"  Companies: {len(registry)}")

    # Preload 10-K chunks and KG structures
    company_data = {}
    for cik, info in registry.items():
        name = info["name"]
        ov = info.get("overview", {})
        mission = ov.get("mission", "")

        # 10-K chunks
        filepath = find_10k_file(cik)
        chunks = []
        if filepath:
            item1 = extract_item1(filepath)
            if len(item1) > 100:
                chunks = chunk_text(item1)

        # KG structure
        kg_struct = parse_ontology_ttl(cik, name)

        company_data[cik] = {
            "name": name,
            "sector": info.get("sector", ""),
            "mission": mission,
            "chunks": chunks,
            "kg_struct": kg_struct,
        }

    print(f"  Companies with chunks: {sum(1 for c in company_data.values() if c['chunks'])}")
    print(f"  Companies with KG: {sum(1 for c in company_data.values() if c['kg_struct']['mission'])}")

    # ── Phase 1: Generate questions ──
    print("\n[Phase 1] Generating questions...")
    questions = generate_questions(registry)
    print(f"  Total questions: {len(questions)}")

    # ── Phase 2: Run all systems (Generator = Qwen-7B) ──
    print("\n[Phase 2] Running 5 reasoning systems + 3 ablation variants...")
    swap_to_model(GENERATOR_MODEL, JUDGE_MODEL)

    SYSTEMS = {
        "Text-Only LLM": system_text_only,
        "Vector Retrieval + LLM": system_vector_retrieval,
        "SentenceTransformer + RAG": system_rag,
        "GraphRAG": system_graphrag,
        "MSG-KG Reasoning": system_msgkg,
    }

    ABLATIONS = {
        "Full MSG-KG Framework": system_msgkg,
        "Without Graph Retrieval": ablation_no_graph,
        "Without Evidence Linking": ablation_no_evidence,
        "Without KG Structure": ablation_no_kg_structure,
    }

    ALL_SYSTEMS = {**SYSTEMS, **ABLATIONS}
    # "Full MSG-KG Framework" and "MSG-KG Reasoning" are same function — deduplicate
    all_answers = {}  # {qid: {system_name: answer_text}}

    checkpoint_path = RESULTS_DIR / "checkpoint_answers.json"
    if checkpoint_path.exists():
        with open(checkpoint_path, encoding="utf-8") as f:
            all_answers = json.load(f)
        print(f"  Resuming from checkpoint ({len(all_answers)} questions done)")

    total_calls = len(questions) * len(ALL_SYSTEMS)
    call_count = 0

    for qi, q in enumerate(questions):
        qid = q["qid"]
        cik = q["cik"]
        cd = company_data[cik]

        if qid in all_answers and len(all_answers[qid]) >= len(ALL_SYSTEMS):
            call_count += len(ALL_SYSTEMS)
            continue

        if qid not in all_answers:
            all_answers[qid] = {}

        for sys_name, sys_func in ALL_SYSTEMS.items():
            if sys_name in all_answers[qid]:
                call_count += 1
                continue

            # Reuse MSG-KG answer for "Full MSG-KG Framework"
            if sys_name == "Full MSG-KG Framework" and "MSG-KG Reasoning" in all_answers[qid]:
                all_answers[qid][sys_name] = all_answers[qid]["MSG-KG Reasoning"]
                call_count += 1
                continue

            call_count += 1
            if call_count % 20 == 0:
                pct = call_count / total_calls * 100
                print(f"  Progress: {call_count}/{total_calls} ({pct:.0f}%) — {q['company'][:20]} / {sys_name}")

            answer = sys_func(
                question=q["question"],
                chunks=cd["chunks"],
                kg_struct=cd["kg_struct"],
                mission=cd["mission"],
            )
            all_answers[qid][sys_name] = answer or ""

        # Checkpoint every 10 questions
        if (qi + 1) % 10 == 0:
            with open(checkpoint_path, "w", encoding="utf-8") as f:
                json.dump(all_answers, f, ensure_ascii=False)

    # Final save
    with open(checkpoint_path, "w", encoding="utf-8") as f:
        json.dump(all_answers, f, ensure_ascii=False)
    print(f"  Generation complete. Total answers: {sum(len(v) for v in all_answers.values())}")

    # ── Phase 3: Judge all answers (Judge = Llama-3.1-8B) ──
    print("\n[Phase 3] Judging with Llama-3.1-8B...")
    swap_to_model(JUDGE_MODEL, GENERATOR_MODEL)

    all_scores = {}  # {qid: {system_name: {"grounding": N, "explanation": N}}}
    scores_checkpoint = RESULTS_DIR / "checkpoint_scores.json"
    if scores_checkpoint.exists():
        with open(scores_checkpoint, encoding="utf-8") as f:
            all_scores = json.load(f)
        print(f"  Resuming from checkpoint ({len(all_scores)} questions judged)")

    total_judge = len(questions) * len(ALL_SYSTEMS)
    judge_count = 0

    for qi, q in enumerate(questions):
        qid = q["qid"]
        cik = q["cik"]
        cd = company_data[cik]

        if qid in all_scores and len(all_scores[qid]) >= len(ALL_SYSTEMS):
            judge_count += len(ALL_SYSTEMS)
            continue

        if qid not in all_scores:
            all_scores[qid] = {}

        for sys_name in ALL_SYSTEMS:
            if sys_name in all_scores[qid]:
                judge_count += 1
                continue

            # Reuse score for "Full MSG-KG Framework"
            if sys_name == "Full MSG-KG Framework" and "MSG-KG Reasoning" in all_scores[qid]:
                all_scores[qid][sys_name] = all_scores[qid]["MSG-KG Reasoning"]
                judge_count += 1
                continue

            answer = all_answers.get(qid, {}).get(sys_name, "")
            if not answer:
                all_scores[qid][sys_name] = {"grounding": 1, "explanation": 1}
                judge_count += 1
                continue

            judge_count += 1
            if judge_count % 20 == 0:
                pct = judge_count / total_judge * 100
                print(f"  Judging: {judge_count}/{total_judge} ({pct:.0f}%) — {q['company'][:20]} / {sys_name}")

            g_score, e_score = judge_answer(q, answer, cd["mission"], cd["kg_struct"])
            all_scores[qid][sys_name] = {"grounding": g_score, "explanation": e_score}

        # Checkpoint every 10
        if (qi + 1) % 10 == 0:
            with open(scores_checkpoint, "w", encoding="utf-8") as f:
                json.dump(all_scores, f, ensure_ascii=False)

    with open(scores_checkpoint, "w", encoding="utf-8") as f:
        json.dump(all_scores, f, ensure_ascii=False)
    print(f"  Judging complete.")

    # ── Phase 4: Compute statistics ──
    print("\n[Phase 4] Computing statistics...")

    def collect_scores(system_name, metric):
        """Collect all scores for a system across questions."""
        scores = []
        for qid in all_scores:
            s = all_scores[qid].get(system_name, {}).get(metric, 2)
            scores.append(s)
        return normalize_score(scores)

    # Table 4: Reasoning results
    print("\n" + "=" * 70)
    print("TABLE 4: Evidence-Grounded Reasoning Performance")
    print("=" * 70)
    print(f"{'Model':<30} {'Grounding Acc':>15} {'Explanation Q':>15}")
    print("-" * 62)

    table4_data = {}
    for sys_name in SYSTEMS:
        g_scores = collect_scores(sys_name, "grounding")
        e_scores = collect_scores(sys_name, "explanation")
        g_mean, g_lo, g_hi = bootstrap_ci(g_scores)
        e_mean, e_lo, e_hi = bootstrap_ci(e_scores)

        table4_data[sys_name] = {
            "grounding_mean": g_mean, "grounding_ci": (g_lo, g_hi),
            "explanation_mean": e_mean, "explanation_ci": (e_lo, e_hi),
            "grounding_scores": g_scores, "explanation_scores": e_scores,
        }

        g_str = f"{g_mean:.1f} ({g_lo:.1f}-{g_hi:.1f})"
        e_str = f"{e_mean:.1f} ({e_lo:.1f}-{e_hi:.1f})"
        bold = "**" if sys_name == "MSG-KG Reasoning" else ""
        print(f"{bold}{sys_name:<30}{bold} {g_str:>15} {e_str:>15}")

    # Significance tests (MSG-KG vs each baseline)
    print("\nSignificance tests (MSG-KG vs baselines, paired bootstrap):")
    msgkg_g = table4_data["MSG-KG Reasoning"]["grounding_scores"]
    msgkg_e = table4_data["MSG-KG Reasoning"]["explanation_scores"]
    for sys_name in SYSTEMS:
        if sys_name == "MSG-KG Reasoning":
            continue
        p_g = paired_bootstrap_test(msgkg_g, table4_data[sys_name]["grounding_scores"])
        p_e = paired_bootstrap_test(msgkg_e, table4_data[sys_name]["explanation_scores"])
        sig_g = "*" if p_g < 0.05 else ""
        sig_e = "*" if p_e < 0.05 else ""
        print(f"  vs {sys_name:<28} Grounding p={p_g:.3f}{sig_g}  Explanation p={p_e:.3f}{sig_e}")

    # Table 5: Ablation
    print("\n" + "=" * 70)
    print("TABLE 5: Ablation Study")
    print("=" * 70)
    print(f"{'Model Variant':<30} {'Grounding Acc':>15}")
    print("-" * 47)

    table5_data = {}
    for variant_name in ABLATIONS:
        g_scores = collect_scores(variant_name, "grounding")
        g_mean, g_lo, g_hi = bootstrap_ci(g_scores)
        table5_data[variant_name] = {
            "grounding_mean": g_mean, "grounding_ci": (g_lo, g_hi),
            "grounding_scores": g_scores,
        }
        g_str = f"{g_mean:.1f} ({g_lo:.1f}-{g_hi:.1f})"
        print(f"{variant_name:<30} {g_str:>15}")

    # ── Phase 5: Human evaluation subset ──
    print("\n[Phase 5] Selecting human evaluation subset...")
    eval_items = []
    for q in questions:
        qid = q["qid"]
        item = {**q, "scores": all_scores.get(qid, {}), "answers": all_answers.get(qid, {})}
        eval_items.append(item)

    human_subset = select_human_eval_subset(eval_items, HUMAN_EVAL_SIZE)
    human_eval_path = RESULTS_DIR / "human_eval_subset.json"
    with open(human_eval_path, "w", encoding="utf-8") as f:
        json.dump(human_subset, f, indent=2, ensure_ascii=False)
    print(f"  Saved {len(human_subset)} examples to {human_eval_path}")

    # ── Phase 6: Save all results ──
    print("\n[Phase 6] Saving results...")

    results = {
        "table4": table4_data,
        "table5": table5_data,
        "config": {
            "generator": GENERATOR_MODEL,
            "judge": JUDGE_MODEL,
            "n_questions": len(questions),
            "n_companies": len(registry),
            "n_bootstrap": N_BOOTSTRAP,
            "confidence": CONFIDENCE,
            "question_types": ["factual", "strategic", "trace"],
        },
    }

    # Clean numpy for JSON serialization
    def clean_for_json(obj):
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, dict):
            return {k: clean_for_json(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [clean_for_json(i) for i in obj]
        return obj

    with open(RESULTS_DIR / "table4_results.json", "w", encoding="utf-8") as f:
        json.dump(clean_for_json(table4_data), f, indent=2)
    with open(RESULTS_DIR / "table5_results.json", "w", encoding="utf-8") as f:
        json.dump(clean_for_json(table5_data), f, indent=2)
    with open(RESULTS_DIR / "all_answers.json", "w", encoding="utf-8") as f:
        json.dump(all_answers, f, ensure_ascii=False)
    with open(RESULTS_DIR / "all_scores.json", "w", encoding="utf-8") as f:
        json.dump(clean_for_json(all_scores), f, indent=2)
    with open(RESULTS_DIR / "full_results.json", "w", encoding="utf-8") as f:
        json.dump(clean_for_json(results), f, indent=2)

    # ── Phase 7: Generate methodology .docx ──
    print("\n[Phase 7] Generating methodology document...")
    generate_methodology_doc(results, table4_data, table5_data)

    print("\n" + "=" * 70)
    print("COMPLETE. All results saved to data/eval_results/")
    print("=" * 70)


# ── Methodology document ────────────────────────────────────────────────

def generate_methodology_doc(results, table4_data, table5_data):
    """Generate .docx with methodology + tables."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        print("  python-docx not available, saving as .txt instead")
        _save_methodology_txt(results, table4_data, table5_data)
        return

    doc = Document()

    # Title
    doc.add_heading('EMNLP Evaluation: Evidence-Grounded Reasoning', level=1)

    # Methodology section
    doc.add_heading('1. Evaluation Setup', level=2)
    doc.add_paragraph(
        f"We evaluate the MSG-KG framework on evidence-grounded reasoning using "
        f"{results['config']['n_companies']} S&P companies with SEC 10-K filings. "
        f"We generate {results['config']['n_questions']} test questions "
        f"({results['config']['n_questions'] // results['config']['n_companies']} per company) "
        f"across three question types: factual grounding, strategic reasoning, and cross-element reasoning."
    )

    doc.add_heading('2. Models', level=2)
    doc.add_paragraph(
        f"Generator Model: {results['config']['generator']} (Qwen2.5-7B-Instruct, local GPU via Ollama)\n"
        f"Judge Model: {results['config']['judge']} (Llama-3.1-8B-Instruct, local GPU via Ollama)\n\n"
        f"We use different model families for generation (Qwen) and evaluation (Llama) to avoid "
        f"self-preference bias documented in Panickssery et al. (2024). The judge model evaluates "
        f"all system outputs using explicit rubrics with 1-5 Likert scoring."
    )

    doc.add_heading('3. Reasoning Systems (Table 4)', level=2)
    systems_desc = [
        ("Text-Only LLM", "The question is passed directly to the LLM with no retrieval context. The model relies entirely on parametric knowledge."),
        ("Vector Retrieval + LLM", "TF-IDF vectorization of 10-K text chunks with cosine similarity retrieval. Top-5 chunks are provided as context to the LLM."),
        ("SentenceTransformer + RAG", "Semantic retrieval using all-MiniLM-L6-v2 sentence embeddings. Top-5 chunks retrieved by cosine similarity are provided to the LLM."),
        ("GraphRAG", "Hybrid system combining knowledge graph triple retrieval from FIBO-aligned ontology KG with top-3 semantic text retrieval."),
        ("MSG-KG Reasoning", "Full pipeline: structured Mission → Strategic Pillar → Operational Mechanism reasoning path from KG, plus evidence spans retrieved from 10-K text, plus KG relation triples."),
    ]
    for name, desc in systems_desc:
        doc.add_paragraph(f"{name}: {desc}")

    doc.add_heading('4. Ablation Variants (Table 5)', level=2)
    ablation_desc = [
        ("Full MSG-KG Framework", "Complete pipeline with all components."),
        ("Without Graph Retrieval", "KG graph traversal removed; text retrieval and mission context only."),
        ("Without Evidence Linking", "Evidence span retrieval removed; KG reasoning path only, no textual evidence."),
        ("Without KG Structure", "Structured Mission→Pillar→Mechanism path removed; flat KG elements + text retrieval."),
    ]
    for name, desc in ablation_desc:
        doc.add_paragraph(f"{name}: {desc}")

    doc.add_heading('5. Evaluation Metrics', level=2)
    doc.add_paragraph(
        "Grounding Accuracy: Measures whether generated answers are supported by specific, verifiable "
        "evidence from the 10-K filing. Scored 1-5, normalized to 0-100.\n\n"
        "Explanation Quality: Measures whether the system traces a coherent "
        "Mission → Strategic Pillar → Operational Mechanism chain. Scored 1-5, normalized to 0-100.\n\n"
        "Following the RAGAS framework (Es et al., EACL 2024) and G-Eval protocol (Liu et al., EMNLP 2023), "
        "we use LLM-as-judge with explicit rubrics and single-criterion evaluation."
    )

    doc.add_heading('6. Statistical Methods', level=2)
    doc.add_paragraph(
        f"We report mean scores with {int(results['config']['confidence']*100)}% bootstrap confidence intervals "
        f"(n={results['config']['n_bootstrap']} resamples). Significance testing uses paired bootstrap "
        f"resampling following Dror et al. (2018). * denotes p < 0.05."
    )

    doc.add_heading('7. Human Evaluation', level=2)
    doc.add_paragraph(
        f"A stratified subset of {HUMAN_EVAL_SIZE} examples is selected for human verification, "
        f"balanced across question types and prioritizing examples with high inter-system score variance. "
        f"Human evaluation validates LLM judge reliability (correlation analysis)."
    )

    # Table 4
    doc.add_heading('Table 4: Evidence-Grounded Reasoning Performance', level=2)
    t4 = doc.add_table(rows=6, cols=3)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Model", "Grounding Accuracy", "Explanation Quality"]
    for i, h in enumerate(headers):
        t4.rows[0].cells[i].text = h

    for ri, sys_name in enumerate(["Text-Only LLM", "Vector Retrieval + LLM", "SentenceTransformer + RAG", "GraphRAG", "MSG-KG Reasoning"]):
        d = table4_data[sys_name]
        t4.rows[ri+1].cells[0].text = sys_name
        t4.rows[ri+1].cells[1].text = f"{d['grounding_mean']:.1f}"
        t4.rows[ri+1].cells[2].text = f"{d['explanation_mean']:.1f}"

    # Table 5
    doc.add_heading('Table 5: Ablation Study', level=2)
    t5 = doc.add_table(rows=5, cols=2)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5.rows[0].cells[0].text = "Model Variant"
    t5.rows[0].cells[1].text = "Grounding Accuracy"

    for ri, variant in enumerate(["Full MSG-KG Framework", "Without Graph Retrieval", "Without Evidence Linking", "Without KG Structure"]):
        d = table5_data[variant]
        t5.rows[ri+1].cells[0].text = variant
        t5.rows[ri+1].cells[1].text = f"{d['grounding_mean']:.1f}"

    # Save
    doc_path = RESULTS_DIR / "evaluation_methodology.docx"
    doc.save(str(doc_path))
    print(f"  Saved methodology to {doc_path}")


def _save_methodology_txt(results, table4_data, table5_data):
    """Fallback: save as plain text."""
    lines = [
        "EMNLP Evaluation: Evidence-Grounded Reasoning",
        "=" * 50,
        f"Generator: {results['config']['generator']}",
        f"Judge: {results['config']['judge']}",
        f"Questions: {results['config']['n_questions']}",
        f"Companies: {results['config']['n_companies']}",
        "",
        "TABLE 4: Reasoning Performance",
        f"{'Model':<30} {'Grounding':>10} {'Explanation':>12}",
    ]
    for sys_name in ["Text-Only LLM", "Vector Retrieval + LLM", "SentenceTransformer + RAG", "GraphRAG", "MSG-KG Reasoning"]:
        d = table4_data[sys_name]
        lines.append(f"{sys_name:<30} {d['grounding_mean']:>10.1f} {d['explanation_mean']:>12.1f}")

    lines.extend(["", "TABLE 5: Ablation", f"{'Variant':<30} {'Grounding':>10}"])
    for v in ["Full MSG-KG Framework", "Without Graph Retrieval", "Without Evidence Linking", "Without KG Structure"]:
        d = table5_data[v]
        lines.append(f"{v:<30} {d['grounding_mean']:>10.1f}")

    txt_path = RESULTS_DIR / "evaluation_methodology.txt"
    txt_path.write_text("\n".join(lines), encoding="utf-8")
    print(f"  Saved methodology to {txt_path}")


if __name__ == "__main__":
    main()
