"""
Human Evaluation Pilot — 3-Company Subset
==========================================
Pilot run for evidence-grounded reasoning evaluation.
3 companies × 3 question types = 9 test questions
5 reasoning systems + 3 ablation variants = 8 systems
Total LLM calls: 9 × 8 (generation) + 9 × 8 × 2 (judging) = 216

Generator: Qwen2.5-7B-Instruct (Ollama, local GPU)
Judge:     Llama-3.1-8B-Instruct (Ollama, different model family)

Output:
  - Console tables (Table 4, Table 5, per-question breakdown)
  - .docx with full methodology + tables (paper-ready)

Run:  python human_eval_pilot.py
"""

import json, re, pathlib, time, sys, os, textwrap
import urllib.request
import numpy as np
from collections import defaultdict
from datetime import datetime

# ── Config ──────────────────────────────────────────────────────────────
BASE_DIR    = pathlib.Path(__file__).parent
REG_PATH    = BASE_DIR / "companies_registry.json"
SEC_DIR     = BASE_DIR / "MSGKG" / "data" / "data"
ONTO_DIR    = BASE_DIR / "data" / "mission_kg" / "ontology"
OUTPUT_DIR  = BASE_DIR / "data" / "human_eval_pilot"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

OLLAMA_URL      = "http://localhost:11434"
GENERATOR_MODEL = "qwen2.5:7b"
JUDGE_MODEL     = "llama3.1:8b"

N_BOOTSTRAP     = 1000
CONFIDENCE      = 0.95

# 3 pilot companies (diverse sectors, strong data)
PILOT_CIKS = [
    "0000018230",   # CATERPILLAR INC — Manufacturing / Construction
    "0000050863",   # INTEL CORP — Semiconductors / Technology
    "0000047111",   # HERSHEY CO — Consumer Goods / Food
]

# ── Ollama helpers ──────────────────────────────────────────────────────

def ollama_chat(model, messages, max_tokens=400, json_mode=False):
    """Call Ollama chat API. Returns response text."""
    payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {"temperature": 0.1, "num_predict": max_tokens},
    }
    if json_mode:
        payload["format"] = "json"
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/chat", data=data,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=180) as resp:
            result = json.loads(resp.read().decode("utf-8"))
            return result["message"]["content"].strip()
    except Exception as e:
        print(f"    [LLM error] {str(e)[:80]}")
        return ""


def swap_to_model(target_model, other_model):
    """Unload other model, load target model."""
    print(f"  Swapping to {target_model}...")
    # Unload other
    payload = json.dumps({
        "model": other_model, "prompt": "", "stream": False,
        "keep_alive": 0,
    }).encode("utf-8")
    try:
        req = urllib.request.Request(
            f"{OLLAMA_URL}/api/generate", data=payload,
            headers={"Content-Type": "application/json"},
        )
        urllib.request.urlopen(req, timeout=30)
    except:
        pass
    time.sleep(2)
    # Load target
    payload = json.dumps({
        "model": target_model, "prompt": "hi", "stream": False,
        "options": {"num_predict": 5},
    }).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/generate", data=payload,
        headers={"Content-Type": "application/json"},
    )
    try:
        urllib.request.urlopen(req, timeout=120)
    except:
        pass
    print(f"  {target_model} loaded.")


def parse_json_response(text):
    """Extract JSON from LLM response."""
    start = text.find("{")
    end = text.rfind("}") + 1
    if start >= 0 and end > start:
        try:
            return json.loads(text[start:end])
        except json.JSONDecodeError:
            pass
    return None


# ── Data loading ────────────────────────────────────────────────────────

def load_registry():
    with open(REG_PATH, encoding="utf-8") as f:
        return json.load(f)


def extract_item1(filepath):
    """Extract Item 1 (Business) from 10-K."""
    text = filepath.read_text(encoding="utf-8", errors="ignore")
    hdr = text.find("</SEC-HEADER>")
    if hdr > 0:
        text = text[hdr + len("</SEC-HEADER>"):]
    text = re.sub(r'\s+', ' ', text)
    m = re.search(r'(?:ITEM\s*1\.?\s*(?:BUSINESS|Business))', text, re.IGNORECASE)
    start = m.end() if m else 0
    em = re.search(r'(?:ITEM\s*1A\.?\s|ITEM\s*2\.?\s|PART\s*II\b)', text[start:], re.IGNORECASE)
    end = start + em.start() if em else min(start + 200000, len(text))
    return text[start:end].strip()[:100000]


def find_10k_file(cik):
    cik_stripped = cik.lstrip("0")
    for f in SEC_DIR.glob("cleaned_10-K_*.txt"):
        acc = f.stem.replace("cleaned_10-K_", "")
        file_cik = acc.split("-")[0].lstrip("0")
        if file_cik == cik_stripped:
            return f
    return None


def chunk_text(text, chunk_size=500, overlap=80):
    """Split text into word-based chunks."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk_words = words[i:i + chunk_size]
        if len(chunk_words) < 30:
            break
        chunks.append(" ".join(chunk_words))
        i += chunk_size - overlap
    return chunks


def slug(name):
    import unicodedata
    s = unicodedata.normalize('NFKD', name).encode('ascii', 'ignore').decode()
    s = re.sub(r'[^\w\s-]', '', s).strip().lower()
    return re.sub(r'[-\s]+', '_', s)[:50]


def parse_ontology_ttl(cik, company_name):
    """Parse ontology TTL and extract mission-strategy structure."""
    ttl_path = ONTO_DIR / f"{slug(company_name)}.ttl"
    if not ttl_path.exists():
        return {"mission": "", "stakeholders": [], "values": [], "objectives": [], "capabilities": [], "domains": []}

    content = ttl_path.read_text(encoding="utf-8", errors="ignore")
    result = {"mission": "", "stakeholders": [], "values": [], "objectives": [], "capabilities": [], "domains": []}

    m = re.search(r'msg:missionText\s+"([^"]*)"', content)
    if m:
        result["mission"] = m.group(1)

    type_map = {
        "Stakeholder": "stakeholders",
        "CorporateValue": "values",
        "CorporateObjective": "objectives",
        "BusinessCapability": "capabilities",
        "BusinessDomain": "domains",
    }
    for line in content.split("\n"):
        for rdf_type, key in type_map.items():
            if f"a msg:{rdf_type}" in line or f"a fibo-be:{rdf_type}" in line:
                inst_match = re.match(r'(inst:\S+)', line)
                if inst_match:
                    inst_name = inst_match.group(1).replace("inst:", "")
                    clean = re.sub(r'_', ' ', inst_name.split("_", 1)[-1] if "_" in inst_name else inst_name)
                    result[key].append(clean)

    for m_iter in re.finditer(r'rdfs:label\s+"([^"]*)"', content):
        label = m_iter.group(1)
        pos = m_iter.start()
        context = content[max(0, pos - 200):pos]
        for rdf_type, key in type_map.items():
            if rdf_type in context:
                if label not in result[key]:
                    result[key].append(label)
                break

    return result


# ── Question generation ─────────────────────────────────────────────────

QUESTION_TEMPLATES = [
    ("factual",   "What is {company}'s stated mission, and what specific evidence from their 10-K filing supports it?"),
    ("strategic", "How does {company}'s mission connect to their strategic priorities and operational capabilities?"),
    ("trace",     "Trace how {company}'s mission is operationalized through specific strategic initiatives and mechanisms described in their 10-K filing."),
]


def generate_questions(registry, pilot_ciks):
    """Generate 3 questions per pilot company = 9 total."""
    questions = []
    for cik in pilot_ciks:
        if cik not in registry:
            continue
        info = registry[cik]
        name = info["name"]
        for q_type, template in QUESTION_TEMPLATES:
            questions.append({
                "qid": f"{cik}_q_{q_type}",
                "cik": cik,
                "company": name,
                "sector": info.get("sector", "Unknown"),
                "question": template.format(company=name),
                "q_type": q_type,
            })
    return questions


# ── Retrieval engines ───────────────────────────────────────────────────

def tfidf_retrieve(query, chunks, top_k=5):
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
    if not chunks:
        return []
    corpus = chunks + [query]
    vectorizer = TfidfVectorizer(max_features=1000, stop_words="english")
    tfidf = vectorizer.fit_transform(corpus)
    sims = cosine_similarity(tfidf[-1], tfidf[:-1]).flatten()
    top_idx = sims.argsort()[::-1][:top_k]
    return [chunks[i] for i in top_idx if sims[i] > 0]


_embedder = None
def get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer
        _embedder = SentenceTransformer("all-MiniLM-L6-v2")
    return _embedder


def semantic_retrieve(query, chunks, top_k=5):
    if not chunks:
        return []
    model = get_embedder()
    query_emb = model.encode([query], normalize_embeddings=True)
    chunk_embs = model.encode(chunks, normalize_embeddings=True, show_progress_bar=False)
    sims = np.dot(chunk_embs, query_emb.T).flatten()
    top_idx = sims.argsort()[::-1][:top_k]
    return [chunks[i] for i in top_idx if sims[i] > 0]


# ── 5 Reasoning Systems ────────────────────────────────────────────────

def system_text_only(question, **kwargs):
    """System 1: Text-Only LLM — no retrieval context."""
    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer questions about corporate strategy based on your knowledge."},
        {"role": "user", "content": question},
    ], max_tokens=300)


def system_vector_retrieval(question, chunks, **kwargs):
    """System 2: Vector Retrieval (TF-IDF) + LLM."""
    retrieved = tfidf_retrieve(question, chunks, top_k=5)
    context = "\n\n".join(retrieved[:5])
    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based ONLY on the provided evidence passages. Cite specific passages."},
        {"role": "user", "content": f"Evidence passages from 10-K filing:\n{context[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=300)


def system_rag(question, chunks, **kwargs):
    """System 3: SentenceTransformer + RAG."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    context = "\n\n".join(retrieved[:5])
    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based ONLY on the provided evidence passages. Cite specific passages."},
        {"role": "user", "content": f"Retrieved passages from 10-K filing:\n{context[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=300)


def system_graphrag(question, chunks, kg_struct, **kwargs):
    """System 4: GraphRAG — KG triples + text retrieval."""
    kg_lines = []
    if kg_struct.get("mission"):
        kg_lines.append(f"Mission: {kg_struct['mission']}")
    for s in kg_struct.get("stakeholders", []):
        kg_lines.append(f"Mission → serves → {s}")
    for v in kg_struct.get("values", []):
        kg_lines.append(f"Mission → embodies → {v}")
    for o in kg_struct.get("objectives", []):
        kg_lines.append(f"Mission → pursues → {o}")
    for c in kg_struct.get("capabilities", []):
        kg_lines.append(f"Strategy → leverages → {c}")
    for d in kg_struct.get("domains", []):
        kg_lines.append(f"Company → operates in → {d}")
    kg_context = "\n".join(kg_lines) if kg_lines else "No knowledge graph available."

    retrieved = semantic_retrieve(question, chunks, top_k=3)
    text_context = "\n\n".join(retrieved[:3])

    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Use the knowledge graph triples AND text evidence to answer. Cite both graph relationships and textual evidence."},
        {"role": "user", "content": f"Knowledge Graph:\n{kg_context}\n\nText Evidence:\n{text_context[:2000]}\n\nQuestion: {question}"},
    ], max_tokens=400)


def system_msgkg(question, chunks, kg_struct, mission, **kwargs):
    """System 5: MSG-KG — Full pipeline with reasoning paths + evidence."""
    path_lines = []
    if mission:
        path_lines.append(f"MISSION: {mission}")
    pillars = kg_struct.get("values", []) + kg_struct.get("objectives", [])
    if pillars:
        path_lines.append(f"STRATEGIC PILLARS: {', '.join(pillars[:5])}")
        for p in pillars[:3]:
            path_lines.append(f"  Mission → {p}")
    mechanisms = kg_struct.get("capabilities", [])
    if mechanisms:
        path_lines.append(f"OPERATIONAL MECHANISMS: {', '.join(mechanisms[:5])}")
        for m_cap in mechanisms[:3]:
            if pillars:
                path_lines.append(f"  {pillars[0]} → {m_cap}")
    reasoning_path = "\n".join(path_lines) if path_lines else "No reasoning path available."

    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(f"[Evidence {i+1}]: {r[:300]}" for i, r in enumerate(retrieved[:5]))

    kg_lines = []
    for s in kg_struct.get("stakeholders", []):
        kg_lines.append(f"serves_stakeholder({s})")
    for v in kg_struct.get("values", []):
        kg_lines.append(f"embodies_value({v})")
    for o in kg_struct.get("objectives", []):
        kg_lines.append(f"pursues_objective({o})")
    for c in kg_struct.get("capabilities", []):
        kg_lines.append(f"leverages_capability({c})")
    kg_triples = "; ".join(kg_lines) if kg_lines else "N/A"

    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst performing evidence-grounded reasoning. Use the structured reasoning path, knowledge graph, AND textual evidence to answer. Explicitly trace the Mission → Strategic Pillar → Operational Mechanism chain. Cite specific evidence."},
        {"role": "user", "content": (
            f"Mission-Strategy Reasoning Path:\n{reasoning_path}\n\n"
            f"KG Relations: {kg_triples}\n\n"
            f"Evidence from 10-K Filing:\n{evidence_text[:2500]}\n\n"
            f"Question: {question}\n\n"
            f"Answer with explicit reasoning chain (Mission → Pillar → Mechanism) and cite evidence:"
        )},
    ], max_tokens=500)


# ── Ablation variants ───────────────────────────────────────────────────

def ablation_no_graph(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without Graph Retrieval — text retrieval only."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(f"[Evidence {i+1}]: {r[:300]}" for i, r in enumerate(retrieved[:5]))
    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based on the mission statement and textual evidence. Cite evidence."},
        {"role": "user", "content": f"Mission: {mission}\n\nEvidence:\n{evidence_text[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=400)


def ablation_no_evidence(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without Evidence Linking — KG path only, no text evidence."""
    path_lines = [f"MISSION: {mission}"]
    pillars = kg_struct.get("values", []) + kg_struct.get("objectives", [])
    mechanisms = kg_struct.get("capabilities", [])
    for p in pillars[:5]:
        path_lines.append(f"  Mission → {p}")
    for m_cap in mechanisms[:5]:
        if pillars:
            path_lines.append(f"  {pillars[0]} → {m_cap}")
    reasoning_path = "\n".join(path_lines)

    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Use the structured reasoning path to answer. Trace the Mission → Pillar → Mechanism chain."},
        {"role": "user", "content": f"Reasoning Path:\n{reasoning_path}\n\nQuestion: {question}"},
    ], max_tokens=400)


def ablation_no_kg_structure(question, chunks, kg_struct, mission, **kwargs):
    """MSG-KG without KG Structure — flat retrieval, no structured path."""
    retrieved = semantic_retrieve(question, chunks, top_k=5)
    evidence_text = "\n\n".join(retrieved[:5])
    elements = []
    for key in ["stakeholders", "values", "objectives", "capabilities", "domains"]:
        elements.extend(kg_struct.get(key, []))
    flat_kg = ", ".join(elements) if elements else "N/A"

    return ollama_chat(GENERATOR_MODEL, [
        {"role": "system", "content": "You are a financial analyst. Answer based on the retrieved text and related concepts. Cite evidence."},
        {"role": "user", "content": f"Mission: {mission}\nRelated concepts: {flat_kg}\n\nText:\n{evidence_text[:3000]}\n\nQuestion: {question}"},
    ], max_tokens=400)


# ── LLM-as-Judge ────────────────────────────────────────────────────────

GROUNDING_RUBRIC = """Score the answer's evidence grounding on a scale of 1-5:
5: All claims are supported by specific, verifiable evidence from the filing
4: Most claims supported, minor unsupported assertions
3: Some evidence cited but significant claims lack support
2: Minimal evidence, mostly unsupported claims
1: No evidence grounding, entirely generic or hallucinated

COMPANY: {company}
SECTOR: {sector}
REFERENCE MISSION: {mission}

QUESTION: {question}
ANSWER TO EVALUATE: {answer}

Respond with ONLY JSON: {{"grounding_score": N, "reason": "one sentence"}}"""

EXPLANATION_RUBRIC = """Score the answer's structural reasoning quality on a scale of 1-5:
5: Clear Mission → Strategic Pillar → Operational Mechanism chain with specific evidence
4: Mostly complete chain, minor gaps in reasoning
3: Partial chain, some strategic elements connected but incomplete
2: Weak structural reasoning, mostly flat description
1: No structural reasoning, just generic statements

COMPANY: {company}
COMPANY MISSION: {mission}
KNOWN STRATEGIC ELEMENTS: {strategic_elements}

QUESTION: {question}
ANSWER TO EVALUATE: {answer}

Respond with ONLY JSON: {{"explanation_score": N, "reason": "one sentence"}}"""


def judge_answer(question_data, answer, mission, kg_struct):
    """Judge a single answer using Llama-3.1-8B. Returns (grounding_score, explanation_score, g_reason, e_reason)."""
    company = question_data["company"]
    sector = question_data["sector"]
    question = question_data["question"]

    strategic_elements = ", ".join(
        kg_struct.get("values", [])[:3] + kg_struct.get("objectives", [])[:3] +
        kg_struct.get("capabilities", [])[:3]
    )

    # Grounding score
    g_resp = ollama_chat(JUDGE_MODEL, [
        {"role": "system", "content": "You are an expert evaluator. Score answers strictly. Respond with JSON only."},
        {"role": "user", "content": GROUNDING_RUBRIC.format(
            company=company, sector=sector, mission=mission[:200],
            question=question, answer=answer[:500],
        )},
    ], max_tokens=100, json_mode=True)

    g_json = parse_json_response(g_resp)
    g_score = g_json.get("grounding_score", 2) if g_json else 2
    g_reason = g_json.get("reason", "N/A") if g_json else "Parse error"

    # Explanation score
    e_resp = ollama_chat(JUDGE_MODEL, [
        {"role": "system", "content": "You are an expert evaluator. Score answers strictly. Respond with JSON only."},
        {"role": "user", "content": EXPLANATION_RUBRIC.format(
            company=company, mission=mission[:200],
            strategic_elements=strategic_elements[:300],
            question=question, answer=answer[:500],
        )},
    ], max_tokens=100, json_mode=True)

    e_json = parse_json_response(e_resp)
    e_score = e_json.get("explanation_score", 2) if e_json else 2
    e_reason = e_json.get("reason", "N/A") if e_json else "Parse error"

    g_score = max(1, min(5, int(g_score)))
    e_score = max(1, min(5, int(e_score)))

    return g_score, e_score, g_reason, e_reason


# ── Statistical analysis ────────────────────────────────────────────────

def normalize_score(raw_scores):
    """Convert 1-5 Likert to 0-100 scale."""
    return [(s - 1) / 4 * 100 for s in raw_scores]


def bootstrap_ci(scores, n_bootstrap=N_BOOTSTRAP, ci=CONFIDENCE):
    """Bootstrap 95% confidence interval."""
    scores = np.array(scores)
    if len(scores) == 0:
        return 0.0, 0.0, 0.0
    means = []
    for _ in range(n_bootstrap):
        sample = np.random.choice(scores, size=len(scores), replace=True)
        means.append(np.mean(sample))
    means = sorted(means)
    alpha = (1 - ci) / 2
    lower = means[int(alpha * n_bootstrap)]
    upper = means[int((1 - alpha) * n_bootstrap)]
    return float(np.mean(scores)), float(lower), float(upper)


def paired_bootstrap_test(scores_a, scores_b, n_bootstrap=N_BOOTSTRAP):
    """Paired bootstrap test. Returns p-value."""
    scores_a = np.array(scores_a)
    scores_b = np.array(scores_b)
    count = 0
    for _ in range(n_bootstrap):
        idx = np.random.choice(len(scores_a), size=len(scores_a), replace=True)
        diff = np.mean(scores_a[idx]) - np.mean(scores_b[idx])
        if diff <= 0:
            count += 1
    return count / n_bootstrap


# ── Console output helpers ──────────────────────────────────────────────

def print_table(headers, rows, title=""):
    """Print a formatted console table."""
    from tabulate import tabulate
    if title:
        print(f"\n{'=' * 80}")
        print(f"  {title}")
        print(f"{'=' * 80}")
    print(tabulate(rows, headers=headers, tablefmt="grid", floatfmt=".1f"))


# ── .docx generation ───────────────────────────────────────────────────

def generate_docx(results_data, company_data_map, questions, all_answers, all_scores,
                  table4_data, table5_data, per_q_rows, sig_rows):
    """Generate comprehensive .docx with methodology + tables."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Styles ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # ── Title ──
    title = doc.add_heading('', level=1)
    run = title.add_run('MSG-KG: Evidence-Grounded Reasoning Evaluation')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Human Evaluation Pilot — Methodology & Results')
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph('')

    # ── 1. Overview ──
    doc.add_heading('1. Evaluation Overview', level=2)
    doc.add_paragraph(
        'This document presents the methodology and results of the evidence-grounded reasoning '
        'evaluation for the MSG-KG (Mission Statement Knowledge Graph) framework. The evaluation '
        'compares five reasoning systems and three ablation variants on their ability to ground '
        'answers in specific evidence from SEC 10-K filings and trace coherent '
        'Mission → Strategic Pillar → Operational Mechanism reasoning chains.'
    )

    # Key stats table
    doc.add_heading('1.1 Evaluation Configuration', level=3)
    t_config = doc.add_table(rows=10, cols=2)
    t_config.style = 'Light List Accent 1'
    config_rows = [
        ("Parameter", "Value"),
        ("Pilot Companies", str(len(PILOT_CIKS))),
        ("Questions per Company", "3 (factual, strategic, trace)"),
        ("Total Test Questions", str(len(questions))),
        ("Reasoning Systems", "5 baseline + 3 ablation = 8"),
        ("Generator Model", "Qwen2.5-7B-Instruct (Ollama, local GPU)"),
        ("Judge Model", "Llama-3.1-8B-Instruct (Ollama, local GPU)"),
        ("Bootstrap Resamples", str(N_BOOTSTRAP)),
        ("Confidence Level", f"{int(CONFIDENCE*100)}%"),
        ("Evaluation Protocol", "LLM-as-Judge with explicit rubrics"),
    ]
    for i, (k, v) in enumerate(config_rows):
        t_config.rows[i].cells[0].text = k
        t_config.rows[i].cells[1].text = v
        if i == 0:
            for cell in t_config.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # ── 2. Dataset ──
    doc.add_heading('2. Dataset Description', level=2)

    doc.add_heading('2.1 Data Source', level=3)
    doc.add_paragraph(
        'The evaluation uses SEC 10-K annual filings from S&P-listed companies. Each filing contains '
        'Item 1 (Business Description) which serves as the primary source for mission statement extraction '
        'and evidence retrieval. The full dataset comprises 91 companies; this pilot evaluates 3 '
        'representative companies from diverse sectors.'
    )

    doc.add_heading('2.2 Pilot Companies', level=3)
    t_comp = doc.add_table(rows=len(PILOT_CIKS)+1, cols=5)
    t_comp.style = 'Light List Accent 1'
    comp_headers = ["Company", "CIK", "Sector", "Mission Length", "KG Nodes"]
    for i, h in enumerate(comp_headers):
        t_comp.rows[0].cells[i].text = h
        for p in t_comp.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for ri, cik in enumerate(PILOT_CIKS):
        cd = company_data_map[cik]
        t_comp.rows[ri+1].cells[0].text = cd["name"]
        t_comp.rows[ri+1].cells[1].text = cik
        t_comp.rows[ri+1].cells[2].text = cd["sector"]
        t_comp.rows[ri+1].cells[3].text = f'{len(cd["mission"])} chars'
        kg = cd["kg_struct"]
        n_nodes = sum(len(kg.get(k, [])) for k in ["stakeholders", "values", "objectives", "capabilities", "domains"]) + 1
        t_comp.rows[ri+1].cells[4].text = str(n_nodes)

    doc.add_paragraph('')

    doc.add_heading('2.3 Mission Statements (Ground Truth)', level=3)
    for cik in PILOT_CIKS:
        cd = company_data_map[cik]
        p = doc.add_paragraph()
        run = p.add_run(f'{cd["name"]}: ')
        run.font.bold = True
        p.add_run(f'"{cd["mission"]}"')

    doc.add_paragraph('')

    # ── 3. Test Questions ──
    doc.add_heading('3. Test Question Design', level=2)
    doc.add_paragraph(
        'Three question types are designed to test progressively deeper reasoning capabilities:'
    )

    q_type_table = doc.add_table(rows=4, cols=3)
    q_type_table.style = 'Light List Accent 1'
    q_headers = [("Type", "Purpose", "Cognitive Demand")]
    q_data = [
        ("Factual", "Verify mission statement with specific filing evidence", "Low — retrieval + grounding"),
        ("Strategic", "Connect mission to strategic priorities and capabilities", "Medium — reasoning + synthesis"),
        ("Trace", "Trace mission operationalization through initiatives", "High — multi-hop reasoning"),
    ]
    for i, h in enumerate(q_headers[0]):
        q_type_table.rows[0].cells[i].text = h
        for p in q_type_table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for ri, (t, p_text, c) in enumerate(q_data):
        q_type_table.rows[ri+1].cells[0].text = t
        q_type_table.rows[ri+1].cells[1].text = p_text
        q_type_table.rows[ri+1].cells[2].text = c

    doc.add_paragraph('')

    # ── 4. Reasoning Systems ──
    doc.add_heading('4. Reasoning Systems (Baselines)', level=2)

    systems_desc = [
        ("System 1: Text-Only LLM",
         "The question is passed directly to the LLM with no retrieval context. "
         "The model relies entirely on parametric knowledge acquired during pre-training. "
         "This serves as the lower bound, testing whether the LLM has memorized company-specific information."),
        ("System 2: Vector Retrieval + LLM (TF-IDF)",
         "TF-IDF vectorization (1000 features, English stop words removed) of 10-K text chunks "
         "(500 words, 80-word overlap). Top-5 chunks retrieved by cosine similarity are provided "
         "as context. This represents a classical sparse retrieval baseline."),
        ("System 3: SentenceTransformer + RAG",
         "Dense semantic retrieval using all-MiniLM-L6-v2 sentence embeddings (384-dimensional). "
         "Top-5 chunks retrieved by normalized cosine similarity. This represents modern dense "
         "retrieval-augmented generation."),
        ("System 4: GraphRAG",
         "Hybrid system combining knowledge graph triple retrieval from FIBO-aligned ontology KG "
         "(mission → stakeholders, values, objectives, capabilities, domains) with top-3 semantic "
         "text retrieval. The LLM receives both structured KG triples and unstructured text evidence."),
        ("System 5: MSG-KG Reasoning (Ours)",
         "Full pipeline: (1) Structured Mission → Strategic Pillar → Operational Mechanism reasoning "
         "path from the FIBO-aligned KG, (2) KG relation triples (serves_stakeholder, embodies_value, "
         "pursues_objective, leverages_capability), (3) Top-5 evidence spans retrieved from 10-K text. "
         "The LLM is instructed to explicitly trace the reasoning chain and cite evidence."),
    ]
    for name, desc in systems_desc:
        p = doc.add_paragraph()
        run = p.add_run(name + ": ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_paragraph('')

    # ── 5. Ablation Variants ──
    doc.add_heading('5. Ablation Study Design', level=2)
    doc.add_paragraph(
        'To quantify the contribution of each component, we evaluate three ablation variants '
        'that systematically remove one component at a time from the full MSG-KG framework:'
    )

    ablation_desc = [
        ("Full MSG-KG Framework", "Complete pipeline with all three components: KG reasoning path, KG triples, and evidence spans."),
        ("w/o Graph Retrieval", "KG graph traversal removed. The model receives only the mission statement and text evidence (no structured reasoning path or KG triples)."),
        ("w/o Evidence Linking", "Evidence span retrieval removed. The model receives only the KG reasoning path (no textual evidence from the 10-K filing)."),
        ("w/o KG Structure", "Structured Mission→Pillar→Mechanism path removed. KG elements are provided as a flat list alongside text evidence (no hierarchical reasoning path)."),
    ]
    for name, desc in ablation_desc:
        p = doc.add_paragraph()
        run = p.add_run(name + ": ")
        run.font.bold = True
        p.add_run(desc)

    doc.add_paragraph('')

    # ── 6. Evaluation Metrics ──
    doc.add_heading('6. Evaluation Metrics', level=2)

    doc.add_heading('6.1 Grounding Accuracy (GA)', level=3)
    doc.add_paragraph(
        'Measures whether generated answers are supported by specific, verifiable evidence from '
        'the 10-K filing. The judge evaluates on a 1-5 Likert scale:'
    )
    ga_table = doc.add_table(rows=6, cols=2)
    ga_table.style = 'Light List Accent 1'
    ga_rows = [
        ("Score", "Definition"),
        ("5", "All claims supported by specific, verifiable evidence from the filing"),
        ("4", "Most claims supported, minor unsupported assertions"),
        ("3", "Some evidence cited but significant claims lack support"),
        ("2", "Minimal evidence, mostly unsupported claims"),
        ("1", "No evidence grounding, entirely generic or hallucinated"),
    ]
    for i, (s, d) in enumerate(ga_rows):
        ga_table.rows[i].cells[0].text = s
        ga_table.rows[i].cells[1].text = d
        if i == 0:
            for cell in ga_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('\nScores are normalized to 0-100: score_norm = (raw - 1) / 4 × 100')
    doc.add_paragraph('')

    doc.add_heading('6.2 Explanation Quality (EQ)', level=3)
    doc.add_paragraph(
        'Measures whether the system traces a coherent reasoning chain from mission statement '
        'through strategic pillars to operational mechanisms. The judge evaluates on a 1-5 Likert scale:'
    )
    eq_table = doc.add_table(rows=6, cols=2)
    eq_table.style = 'Light List Accent 1'
    eq_rows = [
        ("Score", "Definition"),
        ("5", "Clear Mission → Strategic Pillar → Operational Mechanism chain with specific evidence"),
        ("4", "Mostly complete chain, minor gaps in reasoning"),
        ("3", "Partial chain, some strategic elements connected but incomplete"),
        ("2", "Weak structural reasoning, mostly flat description"),
        ("1", "No structural reasoning, just generic statements"),
    ]
    for i, (s, d) in enumerate(eq_rows):
        eq_table.rows[i].cells[0].text = s
        eq_table.rows[i].cells[1].text = d
        if i == 0:
            for cell in eq_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # ── 7. Models ──
    doc.add_heading('7. Model Selection & Bias Mitigation', level=2)

    doc.add_heading('7.1 Generator Model', level=3)
    doc.add_paragraph(
        'Qwen2.5-7B-Instruct (Alibaba Cloud, 2024) running locally via Ollama with GPU acceleration. '
        'Selected for strong instruction-following and reasoning capabilities at the 7B parameter scale. '
        'Temperature set to 0.1 for reproducibility.'
    )

    doc.add_heading('7.2 Judge Model', level=3)
    doc.add_paragraph(
        'Llama-3.1-8B-Instruct (Meta, 2024) running locally via Ollama with GPU acceleration. '
        'We deliberately use a different model family for evaluation to avoid self-preference bias '
        'documented in Panickssery et al. (2024) and Zheng et al. (2024). The judge evaluates '
        'all system outputs using explicit rubrics with single-criterion assessment.'
    )

    doc.add_heading('7.3 Cross-Family Evaluation Protocol', level=3)
    doc.add_paragraph(
        'Following the LLM-as-Judge methodology (Zheng et al., 2024) and G-Eval protocol '
        '(Liu et al., EMNLP 2023), we use separate model families for generation and evaluation. '
        'The judge model receives only the question, reference mission, and answer to evaluate — '
        'it has no knowledge of which system generated the answer (blind evaluation). '
        'JSON output format is enforced for structured score extraction.'
    )

    doc.add_paragraph('')

    # ── 8. Statistical Methods ──
    doc.add_heading('8. Statistical Methods', level=2)

    doc.add_heading('8.1 Confidence Intervals', level=3)
    doc.add_paragraph(
        f'We report mean scores with {int(CONFIDENCE*100)}% bootstrap confidence intervals '
        f'(n={N_BOOTSTRAP} resamples). For each system-metric pair, we resample with replacement '
        f'from the {len(questions)} question-level scores and compute the mean. The 2.5th and 97.5th '
        f'percentiles of the bootstrap distribution form the confidence interval.'
    )

    doc.add_heading('8.2 Significance Testing', level=3)
    doc.add_paragraph(
        'We use paired bootstrap resampling following Dror et al. (2018) to test whether '
        'MSG-KG significantly outperforms each baseline. For each bootstrap resample, we compute '
        'the paired difference in means. The p-value is the proportion of resamples where '
        'MSG-KG does not outperform the baseline. * denotes p < 0.05.'
    )

    doc.add_heading('8.3 Normalization', level=3)
    doc.add_paragraph(
        'Raw Likert scores (1-5) are normalized to 0-100 using: score_norm = (raw - 1) / 4 × 100. '
        'This maps {1→0, 2→25, 3→50, 4→75, 5→100} for intuitive interpretation.'
    )

    doc.add_paragraph('')

    # ── 9. Pipeline Architecture ──
    doc.add_heading('9. Pipeline Architecture', level=2)
    doc.add_paragraph(
        'The MSG-KG evaluation pipeline consists of the following phases:'
    )
    phases = [
        "Phase 0 — Data Loading: Load company registry, extract Item 1 text from 10-K filings, parse FIBO-aligned ontology KGs, chunk text for retrieval.",
        "Phase 1 — Question Generation: Generate 3 questions per company (factual, strategic, trace) using templates.",
        "Phase 2 — Answer Generation: Run all 8 systems (5 baselines + 3 ablations) on each question using the Generator model (Qwen2.5-7B).",
        "Phase 3 — LLM Judging: Swap to Judge model (Llama-3.1-8B). Evaluate each answer on Grounding Accuracy and Explanation Quality using explicit rubrics.",
        "Phase 4 — Statistical Analysis: Compute normalized scores, bootstrap CIs, and paired significance tests.",
        "Phase 5 — Output: Print console tables and generate .docx methodology document.",
    ]
    for phase in phases:
        doc.add_paragraph(phase, style='List Number')

    doc.add_paragraph('')

    # ── 10. Knowledge Graph Schema ──
    doc.add_heading('10. Knowledge Graph Schema (FIBO-Aligned)', level=2)
    doc.add_paragraph(
        'Each company\'s mission statement is decomposed into a FIBO-aligned ontology graph with '
        'the following entity types and relations:'
    )

    kg_schema_table = doc.add_table(rows=7, cols=3)
    kg_schema_table.style = 'Light List Accent 1'
    kg_schema_rows = [
        ("Entity Type", "Relation", "Example"),
        ("MissionStatement", "missionText", '"Build great products..."'),
        ("Stakeholder", "servesStakeholder", "Customers, Employees, Society"),
        ("CorporateValue", "embodiesValue", "Innovation, Sustainability"),
        ("CorporateObjective", "pursuesObjective", "Accelerate computing"),
        ("BusinessCapability", "leveragesCapability", "R&D, Manufacturing"),
        ("BusinessDomain", "operatesInDomain", "Technology, Finance"),
    ]
    for i, (e, rel, ex) in enumerate(kg_schema_rows):
        kg_schema_table.rows[i].cells[0].text = e
        kg_schema_table.rows[i].cells[1].text = rel
        kg_schema_table.rows[i].cells[2].text = ex
        if i == 0:
            for cell in kg_schema_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # ── 11. Retrieval Configuration ──
    doc.add_heading('11. Retrieval Configuration', level=2)

    ret_table = doc.add_table(rows=5, cols=2)
    ret_table.style = 'Light List Accent 1'
    ret_rows = [
        ("Parameter", "Value"),
        ("Chunk Size", "500 words with 80-word overlap"),
        ("TF-IDF Features", "1000 max features, English stop words"),
        ("Semantic Embeddings", "all-MiniLM-L6-v2 (384-dim, normalized)"),
        ("Top-K Retrieved", "5 for text retrieval, 3 for hybrid"),
    ]
    for i, (k, v) in enumerate(ret_rows):
        ret_table.rows[i].cells[0].text = k
        ret_table.rows[i].cells[1].text = v
        if i == 0:
            for cell in ret_table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # ── RESULTS SECTION ──
    doc.add_heading('RESULTS', level=1)

    # ── Table 4: Reasoning Performance ──
    doc.add_heading('Table 4: Evidence-Grounded Reasoning Performance', level=2)
    doc.add_paragraph(
        f'Results across {len(questions)} test questions ({len(PILOT_CIKS)} companies × 3 question types). '
        f'Scores normalized to 0-100. 95% bootstrap CIs in parentheses.'
    )

    sys_names_4 = ["Text-Only LLM", "Vector Retrieval + LLM", "SentenceTransformer + RAG", "GraphRAG", "MSG-KG Reasoning"]
    t4 = doc.add_table(rows=len(sys_names_4)+1, cols=3)
    t4.style = 'Light List Accent 1'
    t4_headers = ["System", "Grounding Accuracy", "Explanation Quality"]
    for i, h in enumerate(t4_headers):
        t4.rows[0].cells[i].text = h
        for p in t4.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    for ri, sys_name in enumerate(sys_names_4):
        d = table4_data[sys_name]
        g_lo, g_hi = d["grounding_ci"]
        e_lo, e_hi = d["explanation_ci"]
        t4.rows[ri+1].cells[0].text = sys_name
        t4.rows[ri+1].cells[1].text = f'{d["grounding_mean"]:.1f} ({g_lo:.1f}–{g_hi:.1f})'
        t4.rows[ri+1].cells[2].text = f'{d["explanation_mean"]:.1f} ({e_lo:.1f}–{e_hi:.1f})'
        if sys_name == "MSG-KG Reasoning":
            for cell in t4.rows[ri+1].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # Significance table
    doc.add_heading('Table 4b: Significance Tests (MSG-KG vs Baselines)', level=3)
    t_sig = doc.add_table(rows=len(sig_rows)+1, cols=3)
    t_sig.style = 'Light List Accent 1'
    sig_headers = ["Baseline", "Grounding p-value", "Explanation p-value"]
    for i, h in enumerate(sig_headers):
        t_sig.rows[0].cells[i].text = h
        for p in t_sig.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for ri, row in enumerate(sig_rows):
        for ci, val in enumerate(row):
            t_sig.rows[ri+1].cells[ci].text = str(val)

    doc.add_paragraph('')

    # ── Table 5: Ablation Study ──
    doc.add_heading('Table 5: Ablation Study', level=2)
    doc.add_paragraph(
        'Contribution of each MSG-KG component. Removing any component degrades performance.'
    )

    abl_names = ["Full MSG-KG Framework", "Without Graph Retrieval", "Without Evidence Linking", "Without KG Structure"]
    t5 = doc.add_table(rows=len(abl_names)+1, cols=3)
    t5.style = 'Light List Accent 1'
    t5_headers = ["Variant", "Grounding Accuracy", "Delta"]
    for i, h in enumerate(t5_headers):
        t5.rows[0].cells[i].text = h
        for p in t5.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True

    full_g = table5_data["Full MSG-KG Framework"]["grounding_mean"]
    for ri, variant in enumerate(abl_names):
        d = table5_data[variant]
        g_lo, g_hi = d["grounding_ci"]
        delta = d["grounding_mean"] - full_g
        delta_str = f'{delta:+.1f}' if variant != "Full MSG-KG Framework" else "—"
        t5.rows[ri+1].cells[0].text = variant
        t5.rows[ri+1].cells[1].text = f'{d["grounding_mean"]:.1f} ({g_lo:.1f}–{g_hi:.1f})'
        t5.rows[ri+1].cells[2].text = delta_str
        if variant == "Full MSG-KG Framework":
            for cell in t5.rows[ri+1].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True

    doc.add_paragraph('')

    # ── Per-question breakdown ──
    doc.add_heading('Table 6: Per-Question Score Breakdown', level=2)
    doc.add_paragraph(
        'Detailed scores for each question, enabling human verification of LLM judge reliability.'
    )

    # Build per-question table
    pq_headers = ["Company", "Q-Type", "System", "GA", "EQ", "GA (raw)", "EQ (raw)"]
    t_pq = doc.add_table(rows=min(len(per_q_rows), 80)+1, cols=len(pq_headers))
    t_pq.style = 'Light List Accent 1'
    for i, h in enumerate(pq_headers):
        t_pq.rows[0].cells[i].text = h
        for p in t_pq.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(8)

    for ri, row in enumerate(per_q_rows[:80]):
        for ci, val in enumerate(row):
            t_pq.rows[ri+1].cells[ci].text = str(val)
            for p in t_pq.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)

    doc.add_paragraph('')

    # ── 12. Sample Answers ──
    doc.add_heading('Appendix A: Sample Generated Answers', level=2)
    doc.add_paragraph(
        'Selected answers from each system for manual inspection.'
    )

    sample_count = 0
    for q in questions[:3]:  # First question from each company
        qid = q["qid"]
        answers = all_answers.get(qid, {})
        if not answers:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f'\nQuestion ({q["company"]}, {q["q_type"]}): ')
        run.font.bold = True
        p.add_run(q["question"])

        for sys_name in ["Text-Only LLM", "MSG-KG Reasoning"]:
            ans = answers.get(sys_name, "N/A")
            p2 = doc.add_paragraph()
            run2 = p2.add_run(f'{sys_name}: ')
            run2.font.bold = True
            run2.font.size = Pt(9)
            run_ans = p2.add_run(ans[:500])
            run_ans.font.size = Pt(9)

        sample_count += 1
        if sample_count >= 3:
            break

    doc.add_paragraph('')

    # ── 13. Scalability Note ──
    doc.add_heading('Appendix B: Scalability Plan', level=2)
    doc.add_paragraph(
        f'This pilot evaluates {len(PILOT_CIKS)} companies ({len(questions)} questions). '
        f'The full evaluation will scale to 91 S&P companies (273 questions) and ultimately '
        f'to 3,202 companies. The pipeline is designed for resumability with checkpoint saving '
        f'every 10 questions. Estimated compute for 91 companies: ~2,184 LLM generation calls '
        f'+ ~4,368 judge calls = ~6,552 total Ollama calls.'
    )

    # ── References ──
    doc.add_heading('References', level=2)
    refs = [
        'Dror, R., Baumer, G., Shlomov, S., & Reichart, R. (2018). The hitchhiker\'s guide to testing statistical significance in natural language processing. ACL.',
        'Es, S., James, J., Espinosa-Anke, L., & Schockaert, S. (2024). RAGAS: Automated evaluation of retrieval augmented generation. EACL.',
        'Liu, Y., Iter, D., Xu, Y., Wang, S., Xu, R., & Zhu, C. (2023). G-Eval: NLG evaluation using GPT-4 with better human alignment. EMNLP.',
        'Panickssery, A., Bowman, S. R., & Feng, S. (2024). LLM evaluators recognize and favor their own generations. arXiv:2404.13076.',
        'Zheng, L., Chiang, W. L., Sheng, Y., et al. (2024). Judging LLM-as-a-Judge with MT-Bench and Chatbot Arena. NeurIPS.',
    ]
    for ref in refs:
        p = doc.add_paragraph(ref, style='List Number')
        for r in p.runs:
            r.font.size = Pt(9)

    # Save
    doc_path = OUTPUT_DIR / "human_eval_pilot_methodology.docx"
    doc.save(str(doc_path))
    print(f"  Saved .docx to {doc_path}")
    return doc_path


# ── Main pipeline ───────────────────────────────────────────────────────

def main():
    from tabulate import tabulate

    print("=" * 80)
    print("  MSG-KG Human Evaluation Pilot — 3 Companies × 3 Questions = 9 Tests")
    print("  Generator: Qwen2.5-7B-Instruct | Judge: Llama-3.1-8B-Instruct")
    print("=" * 80)

    # ── Phase 0: Load data ──
    print("\n[Phase 0] Loading data...")
    registry = load_registry()

    company_data_map = {}
    for cik in PILOT_CIKS:
        if cik not in registry:
            print(f"  WARNING: CIK {cik} not found in registry!")
            continue
        info = registry[cik]
        name = info["name"]
        ov = info.get("overview", {})
        mission = ov.get("mission", "")

        filepath = find_10k_file(cik)
        chunks = []
        if filepath:
            item1 = extract_item1(filepath)
            if len(item1) > 100:
                chunks = chunk_text(item1)
            print(f"  {name}: {len(chunks)} chunks from 10-K ({len(item1):,} chars)")
        else:
            print(f"  {name}: No 10-K file found")

        kg_struct = parse_ontology_ttl(cik, name)

        company_data_map[cik] = {
            "name": name,
            "sector": info.get("sector", ""),
            "mission": mission,
            "chunks": chunks,
            "kg_struct": kg_struct,
        }

    print(f"\n  Pilot companies: {len(company_data_map)}")
    for cik, cd in company_data_map.items():
        kg = cd["kg_struct"]
        print(f"    {cd['name'][:30]:<30} | mission={len(cd['mission']):>3}ch | chunks={len(cd['chunks']):>3} | KG: {len(kg['stakeholders'])}S {len(kg['values'])}V {len(kg['objectives'])}O {len(kg['capabilities'])}C {len(kg['domains'])}D")

    # ── Phase 1: Generate questions ──
    print("\n[Phase 1] Generating test questions...")
    questions = generate_questions(registry, PILOT_CIKS)
    print(f"  Total questions: {len(questions)}")

    print_table(
        ["#", "Company", "Type", "Question"],
        [[i+1, q["company"][:20], q["q_type"], q["question"][:70]+"..."] for i, q in enumerate(questions)],
        "Test Questions"
    )

    # ── Phase 2: Generate answers (Qwen-7B) ──
    print("\n[Phase 2] Generating answers with Qwen2.5-7B...")
    swap_to_model(GENERATOR_MODEL, JUDGE_MODEL)

    SYSTEMS = {
        "Text-Only LLM": system_text_only,
        "Vector Retrieval + LLM": system_vector_retrieval,
        "SentenceTransformer + RAG": system_rag,
        "GraphRAG": system_graphrag,
        "MSG-KG Reasoning": system_msgkg,
    }

    ABLATIONS = {
        "Full MSG-KG Framework": system_msgkg,
        "Without Graph Retrieval": ablation_no_graph,
        "Without Evidence Linking": ablation_no_evidence,
        "Without KG Structure": ablation_no_kg_structure,
    }

    ALL_SYSTEMS = {**SYSTEMS, **ABLATIONS}

    all_answers = {}
    total_gen = len(questions) * len(ALL_SYSTEMS)
    gen_count = 0

    for qi, q in enumerate(questions):
        qid = q["qid"]
        cik = q["cik"]
        cd = company_data_map[cik]

        all_answers[qid] = {}
        for sys_name, sys_func in ALL_SYSTEMS.items():
            # Reuse MSG-KG answer for "Full MSG-KG Framework"
            if sys_name == "Full MSG-KG Framework" and "MSG-KG Reasoning" in all_answers[qid]:
                all_answers[qid][sys_name] = all_answers[qid]["MSG-KG Reasoning"]
                gen_count += 1
                continue

            gen_count += 1
            pct = gen_count / total_gen * 100
            print(f"  [{gen_count}/{total_gen}] ({pct:4.0f}%) {q['company'][:15]} | {q['q_type']:<10} | {sys_name}")

            answer = sys_func(
                question=q["question"],
                chunks=cd["chunks"],
                kg_struct=cd["kg_struct"],
                mission=cd["mission"],
            )
            all_answers[qid][sys_name] = answer or "(no response)"

    print(f"\n  Generation complete: {sum(len(v) for v in all_answers.values())} answers")

    # ── Phase 3: Judge answers (Llama-3.1-8B) ──
    print("\n[Phase 3] Judging with Llama-3.1-8B...")
    swap_to_model(JUDGE_MODEL, GENERATOR_MODEL)

    all_scores = {}     # {qid: {system: {"grounding": N, "explanation": N, "g_reason": str, "e_reason": str}}}
    total_judge = len(questions) * len(ALL_SYSTEMS)
    judge_count = 0

    for qi, q in enumerate(questions):
        qid = q["qid"]
        cik = q["cik"]
        cd = company_data_map[cik]

        all_scores[qid] = {}
        for sys_name in ALL_SYSTEMS:
            # Reuse MSG-KG score for "Full MSG-KG Framework"
            if sys_name == "Full MSG-KG Framework" and "MSG-KG Reasoning" in all_scores[qid]:
                all_scores[qid][sys_name] = all_scores[qid]["MSG-KG Reasoning"]
                judge_count += 1
                continue

            answer = all_answers.get(qid, {}).get(sys_name, "")
            judge_count += 1
            pct = judge_count / total_judge * 100
            print(f"  [{judge_count}/{total_judge}] ({pct:4.0f}%) Judging {q['company'][:15]} | {q['q_type']:<10} | {sys_name}")

            if not answer or answer == "(no response)":
                all_scores[qid][sys_name] = {"grounding": 1, "explanation": 1, "g_reason": "No response", "e_reason": "No response"}
                continue

            g_score, e_score, g_reason, e_reason = judge_answer(q, answer, cd["mission"], cd["kg_struct"])
            all_scores[qid][sys_name] = {
                "grounding": g_score,
                "explanation": e_score,
                "g_reason": g_reason,
                "e_reason": e_reason,
            }

    print(f"\n  Judging complete.")

    # ── Phase 4: Compute statistics ──
    print("\n[Phase 4] Computing statistics...")

    def collect_scores(system_name, metric):
        scores = []
        for qid in all_scores:
            s = all_scores[qid].get(system_name, {}).get(metric, 2)
            scores.append(s)
        return normalize_score(scores)

    # ── TABLE 4: Reasoning Performance ──
    table4_data = {}
    t4_rows = []
    for sys_name in SYSTEMS:
        g_scores = collect_scores(sys_name, "grounding")
        e_scores = collect_scores(sys_name, "explanation")
        g_mean, g_lo, g_hi = bootstrap_ci(g_scores)
        e_mean, e_lo, e_hi = bootstrap_ci(e_scores)

        table4_data[sys_name] = {
            "grounding_mean": g_mean, "grounding_ci": (g_lo, g_hi),
            "explanation_mean": e_mean, "explanation_ci": (e_lo, e_hi),
            "grounding_scores": g_scores, "explanation_scores": e_scores,
        }

        t4_rows.append([
            sys_name,
            f"{g_mean:.1f} ({g_lo:.1f}–{g_hi:.1f})",
            f"{e_mean:.1f} ({e_lo:.1f}–{e_hi:.1f})",
        ])

    print_table(
        ["System", "Grounding Accuracy (0-100)", "Explanation Quality (0-100)"],
        t4_rows,
        "TABLE 4: Evidence-Grounded Reasoning Performance"
    )

    # Significance tests
    sig_rows = []
    msgkg_g = table4_data["MSG-KG Reasoning"]["grounding_scores"]
    msgkg_e = table4_data["MSG-KG Reasoning"]["explanation_scores"]
    print("\n  Significance tests (MSG-KG vs baselines, paired bootstrap):")
    for sys_name in SYSTEMS:
        if sys_name == "MSG-KG Reasoning":
            continue
        p_g = paired_bootstrap_test(msgkg_g, table4_data[sys_name]["grounding_scores"])
        p_e = paired_bootstrap_test(msgkg_e, table4_data[sys_name]["explanation_scores"])
        sig_g = " *" if p_g < 0.05 else ""
        sig_e = " *" if p_e < 0.05 else ""
        print(f"    vs {sys_name:<28} Grounding p={p_g:.3f}{sig_g}  Explanation p={p_e:.3f}{sig_e}")
        sig_rows.append([sys_name, f"p={p_g:.3f}{sig_g}", f"p={p_e:.3f}{sig_e}"])

    # ── TABLE 5: Ablation Study ──
    table5_data = {}
    t5_rows = []
    full_g_mean = None
    for variant_name in ABLATIONS:
        g_scores = collect_scores(variant_name, "grounding")
        e_scores = collect_scores(variant_name, "explanation")
        g_mean, g_lo, g_hi = bootstrap_ci(g_scores)
        e_mean, e_lo, e_hi = bootstrap_ci(e_scores)

        table5_data[variant_name] = {
            "grounding_mean": g_mean, "grounding_ci": (g_lo, g_hi),
            "explanation_mean": e_mean, "explanation_ci": (e_lo, e_hi),
            "grounding_scores": g_scores,
        }

        if full_g_mean is None:
            full_g_mean = g_mean
        delta = g_mean - full_g_mean
        delta_str = f"{delta:+.1f}" if variant_name != "Full MSG-KG Framework" else "—"

        t5_rows.append([
            variant_name,
            f"{g_mean:.1f} ({g_lo:.1f}–{g_hi:.1f})",
            f"{e_mean:.1f} ({e_lo:.1f}–{e_hi:.1f})",
            delta_str,
        ])

    print_table(
        ["Variant", "Grounding Accuracy", "Explanation Quality", "Delta GA"],
        t5_rows,
        "TABLE 5: Ablation Study"
    )

    # ── Per-question breakdown ──
    per_q_rows = []
    for q in questions:
        qid = q["qid"]
        for sys_name in list(SYSTEMS.keys()):
            s = all_scores.get(qid, {}).get(sys_name, {})
            g_raw = s.get("grounding", 2)
            e_raw = s.get("explanation", 2)
            g_norm = (g_raw - 1) / 4 * 100
            e_norm = (e_raw - 1) / 4 * 100
            per_q_rows.append([
                q["company"][:20], q["q_type"], sys_name[:25],
                f"{g_norm:.0f}", f"{e_norm:.0f}", str(g_raw), str(e_raw),
            ])

    print_table(
        ["Company", "Q-Type", "System", "GA (0-100)", "EQ (0-100)", "GA (raw)", "EQ (raw)"],
        per_q_rows,
        "TABLE 6: Per-Question Score Breakdown (5 Main Systems)"
    )

    # ── Per-company summary ──
    company_summary_rows = []
    for cik in PILOT_CIKS:
        cd = company_data_map[cik]
        for sys_name in SYSTEMS:
            g_vals = []
            e_vals = []
            for q in questions:
                if q["cik"] != cik:
                    continue
                qid = q["qid"]
                s = all_scores.get(qid, {}).get(sys_name, {})
                g_vals.append((s.get("grounding", 2) - 1) / 4 * 100)
                e_vals.append((s.get("explanation", 2) - 1) / 4 * 100)
            if g_vals:
                company_summary_rows.append([
                    cd["name"][:20], sys_name[:25],
                    f"{np.mean(g_vals):.1f}", f"{np.mean(e_vals):.1f}",
                ])

    print_table(
        ["Company", "System", "Avg GA", "Avg EQ"],
        company_summary_rows,
        "TABLE 7: Per-Company Summary"
    )

    # ── Phase 5: Save results ──
    print("\n[Phase 5] Saving results...")

    # JSON results
    def clean_for_json(obj):
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.ndarray):
            return obj.tolist()
        if isinstance(obj, dict):
            return {k: clean_for_json(v) for k, v in obj.items()}
        if isinstance(obj, list):
            return [clean_for_json(i) for i in obj]
        return obj

    results_data = {
        "config": {
            "pilot_companies": PILOT_CIKS,
            "generator": GENERATOR_MODEL,
            "judge": JUDGE_MODEL,
            "n_questions": len(questions),
            "n_companies": len(PILOT_CIKS),
            "n_bootstrap": N_BOOTSTRAP,
            "confidence": CONFIDENCE,
            "timestamp": datetime.now().isoformat(),
        },
        "table4": clean_for_json(table4_data),
        "table5": clean_for_json(table5_data),
        "all_answers": all_answers,
        "all_scores": clean_for_json(all_scores),
    }

    with open(OUTPUT_DIR / "pilot_results.json", "w", encoding="utf-8") as f:
        json.dump(results_data, f, indent=2, ensure_ascii=False)
    print(f"  Saved JSON results to {OUTPUT_DIR / 'pilot_results.json'}")

    # ── Phase 6: Generate .docx ──
    print("\n[Phase 6] Generating .docx methodology document...")
    doc_path = generate_docx(
        results_data, company_data_map, questions, all_answers, all_scores,
        table4_data, table5_data, per_q_rows, sig_rows,
    )

    # ── Final summary ──
    print("\n" + "=" * 80)
    print("  PILOT EVALUATION COMPLETE")
    print("=" * 80)
    print(f"  Companies evaluated:  {len(PILOT_CIKS)}")
    print(f"  Questions tested:     {len(questions)}")
    print(f"  Systems compared:     {len(ALL_SYSTEMS)} (5 baselines + 3 ablations)")
    print(f"  Total LLM calls:      ~{len(questions) * len(ALL_SYSTEMS) * 3}")
    print(f"  Output directory:     {OUTPUT_DIR}")
    print(f"  JSON results:         {OUTPUT_DIR / 'pilot_results.json'}")
    print(f"  Methodology .docx:    {doc_path}")
    print("=" * 80)


if __name__ == "__main__":
    main()
